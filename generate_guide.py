"""
TYROSign — Kurumsal Dijital İmza Stüdyosu · Kullanıcı Dokümanı
Generates a single-page Word document with Tiryaki brand colors.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors ──
NAVY   = RGBColor(30, 58, 95)
GOLD   = RGBColor(200, 146, 42)
CYAN   = RGBColor(0, 152, 212)
WHITE  = RGBColor(255, 255, 255)
DARK   = RGBColor(30, 41, 59)
GRAY   = RGBColor(100, 116, 139)
LGRAY  = RGBColor(148, 163, 184)
BG     = RGBColor(248, 250, 252)

doc = Document()

# ── Page Setup (A4, narrow margins) ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(0.8)
section.bottom_margin = Cm(0.6)
section.left_margin   = Cm(1.4)
section.right_margin  = Cm(1.4)

# ── Helper Functions ──
def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_run(paragraph, text, size=9, bold=False, color=DARK, font_name='Calibri'):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    return run

def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def remove_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def set_table_borders(table, color="1e3a5f", size="4"):
    """Remove or set table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def gold_divider():
    """Add a thin gold divider line."""
    p = doc.add_paragraph()
    remove_paragraph_spacing(p, before=3, after=3)
    # Use a bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="c8922a"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_step(number, title, bullets):
    """Add a numbered step with title and bullet points."""
    # Step row as table: [number circle] [content]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl)

    # Number column (narrow)
    num_cell = tbl.cell(0, 0)
    num_cell.width = Cm(1.2)
    set_cell_bg(num_cell, "c8922a")
    set_cell_margins(num_cell, top=40, bottom=40, left=0, right=0)
    np = num_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(np, before=1, after=1)
    nr = np.add_run(str(number))
    nr.font.size = Pt(14)
    nr.font.bold = True
    nr.font.color.rgb = WHITE
    nr.font.name = 'Calibri'

    # Content column
    content_cell = tbl.cell(0, 1)
    set_cell_margins(content_cell, top=30, bottom=30, left=140, right=60)

    # Title
    tp = content_cell.paragraphs[0]
    remove_paragraph_spacing(tp, before=0, after=2, line=13)
    tr = tp.add_run(title)
    tr.font.size = Pt(10)
    tr.font.bold = True
    tr.font.color.rgb = NAVY
    tr.font.name = 'Calibri'

    # Bullets
    for bullet in bullets:
        bp = content_cell.add_paragraph()
        remove_paragraph_spacing(bp, before=0, after=1, line=12)

        # Parse bold markers **text**
        parts = bullet.split('**')
        # Add bullet char
        first = True
        for i, part in enumerate(parts):
            if not part:
                continue
            prefix = '  •  ' if first else ''
            first = False if first and part else first
            if i % 2 == 0:
                # Normal text
                add_styled_run(bp, prefix + part, size=8.5, color=DARK)
            else:
                # Bold text
                add_styled_run(bp, prefix + part, size=8.5, bold=True, color=DARK)

    # Small spacer after step
    spacer = doc.add_paragraph()
    remove_paragraph_spacing(spacer, before=1, after=1)


# ═══════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════

# ── HEADER BANNER ──
header_tbl = doc.add_table(rows=1, cols=2)
header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(header_tbl)

# Make table full width
header_tbl.columns[0].width = Cm(12)
header_tbl.columns[1].width = Cm(6.2)

left_cell = header_tbl.cell(0, 0)
right_cell = header_tbl.cell(0, 1)
set_cell_bg(left_cell, "1e3a5f")
set_cell_bg(right_cell, "1e3a5f")
set_cell_margins(left_cell, top=100, bottom=100, left=160, right=80)
set_cell_margins(right_cell, top=100, bottom=100, left=80, right=160)

# Left: App name + subtitle
lp = left_cell.paragraphs[0]
remove_paragraph_spacing(lp, before=0, after=0)
r1 = lp.add_run("TYRO")
r1.font.size = Pt(16)
r1.font.bold = True
r1.font.color.rgb = WHITE
r1.font.name = 'Calibri'
r2 = lp.add_run("Sign")
r2.font.size = Pt(16)
r2.font.bold = True
r2.font.color.rgb = GOLD
r2.font.name = 'Calibri'

lp2 = left_cell.add_paragraph()
remove_paragraph_spacing(lp2, before=1, after=0)
r3 = lp2.add_run("Kurumsal Dijital İmza Stüdyosu — Kullanıcı Dokümanı")
r3.font.size = Pt(8.5)
r3.font.color.rgb = RGBColor(180, 200, 220)
r3.font.name = 'Calibri'

# Right: URL
rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
remove_paragraph_spacing(rp, before=4, after=0)
r4 = rp.add_run("tyrosign.ttech.business")
r4.font.size = Pt(9)
r4.font.bold = True
r4.font.color.rgb = GOLD
r4.font.name = 'Calibri'

# Small spacer
spacer = doc.add_paragraph()
remove_paragraph_spacing(spacer, before=3, after=1)

# ── STEP 1: Outlook Hazırlığı ──
add_step(1, "Outlook Hazırlığı", [
    'Outlook masaüstü uygulamasında sağ üstteki **"Try the new Outlook"** toggle\'ını açın.',
    'Yeni Outlook açıldığında imza ayarlarına tarayıcı üzerinden erişim sağlanır.',
    'Bu adım **yalnızca bir kere** yapılır — zaten yeni Outlook kullanıyorsanız atlayın.',
])

# ── STEP 2: Siteye Giriş ──
add_step(2, "Siteye Giriş Yapın", [
    'Tarayıcınızdan **tyrosign.ttech.business** adresine gidin.',
    'Giriş ekranının altından dil seçimi yapın: **TR** / **EN** / **RU** / **AR**',
    '**"Microsoft ile Giriş Yap"** butonuna tıklayın — şirket e-posta hesabınızla oturum açın.',
    'İlk girişinizde kısa bir **interaktif rehber** sizi karşılayacak ve arayüzü tanıtacak.',
])

# ── STEP 3: Bilgileri Tamamlayın ──
add_step(3, "Bilgilerinizi Tamamlayın", [
    '**Ad, soyad** ve **profil fotoğrafınız** Microsoft hesabınızdan otomatik gelir.',
    '**Ünvanınızı** Türkçe ve İngilizce olarak girin (ör: Satış Müdürü / Sales Manager).',
    '**Şirket** seçin (31 grup şirketi) → şirket logosu otomatik değişir.',
    '**Ofis** seçin (21 ofis lokasyonu) → adres ve sabit telefon otomatik dolar.',
    'Sağ paneldeki **canlı önizlemede** imzanız anlık olarak güncellenir.',
])

# ── STEP 4: Outlook'a Aktarın ──
add_step(4, "İmzanızı Outlook'a Aktarın", [
    'Alttaki araç çubuğundan **"Outlook"** butonuna tıklayın.',
    'İmzanız otomatik kopyalanır ve **Outlook Web imza ayarları** yeni sekmede açılır.',
    'Açılan sayfada **"+ Yeni imza"** tıklayın → isim verin (ör: "Kurumsal İmza").',
    'İmza kutusuna tıklayıp **Ctrl+V** ile yapıştırın.',
    '**"Yeni iletiler için"** ve **"Yanıtlar/iletmeler için"** kutularından imzanızı seçin.',
    '**Kaydet** butonuna basın. Artık tüm e-postalarınızda kurumsal imzanız görünecek.',
])

# ── Gold divider before tools ──
gold_divider()

# ── STEP 5: Ek Araçlar (Table) ──
p5title = doc.add_paragraph()
remove_paragraph_spacing(p5title, before=3, after=4)
add_styled_run(p5title, "   Ek Araçlar", size=10, bold=True, color=NAVY)

tools_tbl = doc.add_table(rows=4, cols=2)
tools_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(tools_tbl)
tools_tbl.columns[0].width = Cm(3.5)
tools_tbl.columns[1].width = Cm(14.7)

# Header row
for i, header in enumerate(["Araç", "Ne Yapar"]):
    cell = tools_tbl.cell(0, i)
    set_cell_bg(cell, "1e3a5f")
    set_cell_margins(cell, top=40, bottom=40, left=100, right=80)
    p = cell.paragraphs[0]
    remove_paragraph_spacing(p, before=0, after=0)
    add_styled_run(p, header, size=8, bold=True, color=WHITE)

# Tool rows
tools = [
    ("QR Kod", "Kişi bilgilerinizden vCard QR kodu oluşturur. Telefon kamerasıyla okutulduğunda karşı tarafın rehberine doğrudan kayıt oluşturur. Kopyala veya İndir ile paylaşabilirsiniz."),
    ("Kartvizit", "Dijital kartvizit oluşturur — profil fotoğrafınız Office hesabınızdan gelir, üstüne tıklayarak değiştirebilirsiniz. İndir / Kopyala / Paylaş / Dijital Kart butonlarıyla dışa aktarın."),
    ("Yöneticime Bildir", "Ünvan veya iletişim bilgisi değişikliğinde yeni imza tasarımınızı yöneticinize e-posta ile bildirin. Alıcı ve içerik düzenlenebilir."),
]

for idx, (tool, desc) in enumerate(tools):
    row = idx + 1
    bg = "f8fafc" if row % 2 == 1 else "ffffff"

    tool_cell = tools_tbl.cell(row, 0)
    desc_cell = tools_tbl.cell(row, 1)
    set_cell_bg(tool_cell, bg)
    set_cell_bg(desc_cell, bg)
    set_cell_margins(tool_cell, top=35, bottom=35, left=100, right=60)
    set_cell_margins(desc_cell, top=35, bottom=35, left=80, right=80)

    # Tool name
    tp = tool_cell.paragraphs[0]
    remove_paragraph_spacing(tp, before=0, after=0)
    add_styled_run(tp, tool, size=8.5, bold=True, color=GOLD)

    # Description
    dp = desc_cell.paragraphs[0]
    remove_paragraph_spacing(dp, before=0, after=0, line=11.5)
    add_styled_run(dp, desc, size=8, color=DARK)

# ── Tip box ──
spacer = doc.add_paragraph()
remove_paragraph_spacing(spacer, before=3, after=1)

tip_tbl = doc.add_table(rows=1, cols=1)
tip_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(tip_tbl)
tip_cell = tip_tbl.cell(0, 0)
set_cell_bg(tip_cell, "f0f7ff")
set_cell_margins(tip_cell, top=50, bottom=50, left=120, right=120)

# Add left border via cell properties
tc = tip_cell._tc
tcPr = tc.get_or_add_tcPr()
tcBorders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>'
    f'  <w:start w:val="single" w:sz="18" w:space="0" w:color="0098d4"/>'
    f'</w:tcBorders>'
)
tcPr.append(tcBorders)

tp = tip_cell.paragraphs[0]
remove_paragraph_spacing(tp, before=0, after=0, line=12)
add_styled_run(tp, "İpucu: ", size=8, bold=True, color=CYAN)
add_styled_run(tp, "İmza oluşturulduktan sonra ", size=8, color=DARK)
add_styled_run(tp, "Ctrl+C", size=8, bold=True, color=DARK)
add_styled_run(tp, " kısayolu ile de imzanızı hızlıca kopyalayabilirsiniz. Ayarlar sekmesinden imza stilini, renkleri ve sosyal medya bağlantılarını özelleştirebilirsiniz.", size=8, color=DARK)

# ── FOOTER ──
spacer = doc.add_paragraph()
remove_paragraph_spacing(spacer, before=5, after=1)

gold_divider()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
remove_paragraph_spacing(footer_p, before=3, after=0)
add_styled_run(footer_p, "Powered by TTECH Business Solutions", size=7.5, color=LGRAY)
add_styled_run(footer_p, "  ·  ", size=7.5, color=LGRAY)
add_styled_run(footer_p, "Destek: IT ekibinize başvurun veya ", size=7.5, color=LGRAY)
add_styled_run(footer_p, "servicedesk@tiryaki.com.tr", size=7.5, bold=True, color=LGRAY)
add_styled_run(footer_p, " adresine yazın", size=7.5, color=LGRAY)

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TYROSign-Kullanici-Dokumani.docx")
doc.save(output_path)
print(f"Document saved: {output_path}")
